import os
import re
import uuid
from collections import defaultdict

import pandas as pd
from flask import Flask, render_template, request, send_file, flash, redirect, url_for
from werkzeug.utils import secure_filename
from docx import Document
from pypdf import PdfReader
from rapidfuzz import fuzz

APP_ROOT = os.path.dirname(os.path.abspath(__file__))
UPLOAD_FOLDER = os.path.join(APP_ROOT, 'uploads')
OUTPUT_FOLDER = os.path.join(APP_ROOT, 'outputs')
ALLOWED_PM = {'.xlsx', '.xls'}
ALLOWED_STAFFING = {'.docx', '.pdf', '.txt'}
MAX_STAFFING_FILES = 10

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

app = Flask(__name__)
app.secret_key = os.environ.get('SECRET_KEY', 'position-master-tool')
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

SCHOOL_ALIASES = {
    'coquillard': 'Coquillard',
    'darden': 'Darden',
    'harrison': 'Harrison',
    'madison': 'Madison',
}

ROLE_PATTERNS = [
    (r'\bpre[\s-]*k\b|prek|preschool|pre school', 'Pre K'),
    (r'\b1st\b|\bfirst\b|\bgrade 1\b|^1$', 'Grade 1'),
    (r'\b2nd\b|\bsecond\b|\bgrade 2\b|^2$', 'Grade 2'),
    (r'\b3rd\b|\bthird\b|\bgrade 3\b|^3$', 'Grade 3'),
    (r'\b4th\b|\bfourth\b|\bgrade 4\b|^4$', 'Grade 4'),
    (r'\b5th\b|\bfifth\b|\bgrade 5\b|^5$', 'Grade 5'),
    (r'\bkindergarten irs\b', 'IRS'),
    (r'\bel irs\b|\birs\b', 'IRS'),
    (r'\bkindergarten\b|\bk\b', 'Kindergarten'),
    (r'cross\s*cat|cross\s*categorical|\bcc\b', 'Cross Cat'),
    (r'life\s*skills|lifeskills', 'Life Skills'),
    (r'autism|\barr\b', 'Autism / ARR'),
    (r'intervention specialist|intervention', 'Intervention'),
    (r'bilingual education specialist|bilingual educ specialist|bilingual resource specialist', 'Bilingual Specialist'),
    (r'\bel\b.*intervention|eld instructional specialist|el intervention', 'EL Intervention'),
    (r'literacy cadre coach|cadre coach|academic coach|coach', 'Coach'),
    (r'\bmusic\b', 'Music'),
    (r'\bart\b', 'Art'),
    (r'\bgym\b|p\.?e\.?|physical ed|physical education', 'PE'),
    (r'science\s*lab|discovery\s*lab', 'Science Lab'),
    (r'makerspace', 'Makerspace'),
    (r'innovation\s*lab', 'Innovation Lab'),
    (r'speech', 'Speech'),
    (r'\bot\b|occupational', 'OT'),
    (r'\bpt\b|physical\s*therap', 'PT'),
    (r'deaf|hard hearing', 'Deaf/Hard Hearing'),
    (r'title\s*i', 'Title I'),
    (r'digital integration specialist|\bdis\b', 'Specialist'),
    (r'specialist teacher|specialist', 'Specialist'),
]

ROLE_ORDER = [
    'Pre K', 'Kindergarten', 'Grade 1', 'Grade 2', 'Grade 3', 'Grade 4', 'Grade 5',
    'Cross Cat', 'Life Skills', 'Autism / ARR', 'Intervention', 'EL Intervention', 'IRS',
    'Bilingual Specialist', 'Title I', 'Music', 'Art', 'PE', 'Science Lab', 'Makerspace',
    'Innovation Lab', 'Coach', 'Speech', 'OT', 'PT', 'Deaf/Hard Hearing', 'Specialist', 'Other'
]

ROLE_FAMILY = {
    'Pre K': 'Classroom', 'Kindergarten': 'Classroom', 'Grade 1': 'Classroom', 'Grade 2': 'Classroom',
    'Grade 3': 'Classroom', 'Grade 4': 'Classroom', 'Grade 5': 'Classroom',
    'Cross Cat': 'Special Ed', 'Life Skills': 'Special Ed', 'Autism / ARR': 'Special Ed',
    'Intervention': 'Support', 'EL Intervention': 'Support', 'IRS': 'Support',
    'Bilingual Specialist': 'Support', 'Title I': 'Support', 'Coach': 'Support', 'Specialist': 'Support',
    'Speech': 'Support', 'OT': 'Support', 'PT': 'Support', 'Deaf/Hard Hearing': 'Support',
    'Music': 'Specials', 'Art': 'Specials', 'PE': 'Specials', 'Science Lab': 'Specials',
    'Makerspace': 'Specials', 'Innovation Lab': 'Specials', 'Other': 'Other'
}


def allowed(filename, allowed_exts):
    return os.path.splitext(filename.lower())[1] in allowed_exts


def normalize_space(s):
    return re.sub(r'\s+', ' ', str(s or '')).strip()


def infer_school(text, filename=''):
    bucket = f'{filename} {text}'.lower()
    for k, v in SCHOOL_ALIASES.items():
        if k in bucket:
            return v
    cleaned = os.path.splitext(os.path.basename(filename))[0]
    return cleaned[:80] or 'Unknown'


def normalize_role(raw):
    s = str(raw or '')
    s = re.sub(r'\([^)]*\)', ' ', s.lower())
    s = re.sub(r'room\s*#?\s*\d+', ' ', s)
    s = s.replace("teacher’s name", ' ').replace("teacher's name", ' ')
    s = normalize_space(s)
    for pat, label in ROLE_PATTERNS:
        if re.search(pat, s):
            return label
    return 'Other'


def role_family(role):
    return ROLE_FAMILY.get(role, 'Other')


def clean_name_text(s):
    s = normalize_space(s)
    s = re.sub(r'room\s*#?\s*\d+', ' ', s, flags=re.I)
    s = re.sub(r'\([^)]*\)', ' ', s)
    s = re.sub(r'\b(mrs|mr|ms|miss|dr)\.?\b', ' ', s, flags=re.I)
    s = re.sub(r'\b(el certified|el rubric|dli|non-certified|non certified|non-cert|bi|building sub|apprentice teacher|non-certified)\b', ' ', s, flags=re.I)
    s = s.replace("Teacher’s Name", ' ').replace("Teacher's Name", ' ')
    s = s.replace('..', '.')
    s = normalize_space(s).strip(' .-')
    return s


def pm_display_name(raw):
    s = normalize_space(raw)
    if ',' in s:
        last, first = s.split(',', 1)
        s = normalize_space(f'{first} {last}')
    return s


def normalize_name_for_match(name):
    s = clean_name_text(name)
    if not s:
        return ''
    if re.search(r'\bopen\b|\bvacant\b', s, flags=re.I):
        return 'OPEN'
    if ',' in s:
        last, first = s.split(',', 1)
        s = normalize_space(f'{first} {last}')
    s = s.replace('’', "'").replace("'", '')
    s = re.sub(r'[^A-Za-z\s-]', ' ', s)
    s = normalize_space(s).lower()
    tokens = [t for t in s.split() if t not in {'jr', 'sr', 'ii', 'iii', 'iv', 'v', 'm', 'r', 'a', 'e', 'k', 's', 'l', 'j', 'c', 'd', 'n'}]
    return ' '.join(tokens)


def first_last_key(norm_name):
    tokens = [t for t in str(norm_name or '').split() if t]
    if len(tokens) >= 2:
        return f'{tokens[0]} {tokens[-1]}'
    return ' '.join(tokens)


def split_names_blob(text):
    parts = re.split(r'[\n;|]+', str(text or ''))
    out = []
    for p in parts:
        p = clean_name_text(p)
        if p:
            out.append(p)
    return out


def looks_like_person(text):
    s = clean_name_text(text)
    if not s or re.search(r'\b(open|vacant)\b', s, flags=re.I):
        return True
    return bool(re.match(r'^[A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,4}$', s))


def extract_allocations_docx(path):
    doc = Document(path)
    school = infer_school('', os.path.basename(path))
    rows = []
    if doc.tables:
        t = doc.tables[0]
        for r in t.rows[2:]:
            cells = [normalize_space(c.text) for c in r.cells]
            if len(cells) < 6:
                continue
            role = normalize_role(cells[0])
            if role == 'Other':
                continue
            alloc_text = cells[5] or cells[4] or cells[2]
            m = re.search(r'(-?\d+)', alloc_text.replace('·', '-'))
            if m:
                rows.append({'school': school, 'role': role, 'allocation_count': int(m.group(1)), 'allocation_source': 'Pre-Allocation'})
    for t in doc.tables:
        if not t.rows:
            continue
        first = ' '.join(c.text for c in t.rows[0].cells[:2]).lower()
        if 'specials teachers' not in first:
            continue
        for r in t.rows[1:]:
            cells = [normalize_space(c.text) for c in r.cells]
            if len(cells) < 4 or not cells[0]:
                continue
            role = normalize_role(cells[0])
            alloc_text = cells[3] or cells[2] or cells[1]
            m = re.search(r'(-?\d+)', alloc_text.replace('·', '-'))
            if role != 'Other' and m:
                rows.append({'school': school, 'role': role, 'allocation_count': int(m.group(1)), 'allocation_source': 'Specials'})
    return rows


def extract_docx_staffing(path):
    doc = Document(path)
    school = infer_school('', os.path.basename(path))
    rows = []
    seen = set()
    in_staff = False
    for t in doc.tables:
        header = ' '.join(c.text for r in t.rows[:2] for c in r.cells[:2]).lower() if t.rows else ''
        if 'teaching staff names' in header:
            in_staff = True
            continue
        if 'slated not to return' in header:
            in_staff = False
        if not in_staff or len(t.columns) < 2:
            continue
        for r in t.rows:
            role_raw = r.cells[0].text if len(r.cells) > 0 else ''
            names_raw = r.cells[1].text if len(r.cells) > 1 else ''
            if not normalize_space(role_raw) or not normalize_space(names_raw):
                continue
            low = (role_raw + ' ' + names_raw).lower()
            if any(x in low for x in ['continue this list', 'teacher’s name', "teacher's name", 'example']):
                continue
            role_clean = normalize_space(role_raw)
            role = normalize_role(role_clean)
            if role == 'Other' and len(role_clean.split()) > 8:
                continue
            # Fix occasional swapped cells like Carmen Wilber / Bilingual Education Specialist
            if role == 'Other' and normalize_role(names_raw) != 'Other' and looks_like_person(role_clean):
                role_clean, names_raw = names_raw, role_raw
                role = normalize_role(role_clean)
            for person in split_names_blob(names_raw):
                norm = normalize_name_for_match(person)
                if not norm and not re.search(r'[A-Za-z]', person):
                    continue
                rec = (school, os.path.basename(path), role_clean, role, person, norm)
                if rec in seen:
                    continue
                seen.add(rec)
                rows.append({
                    'school': school,
                    'source_file': os.path.basename(path),
                    'source_type': 'DOCX',
                    'raw_role': role_clean,
                    'role': role,
                    'role_family': role_family(role),
                    'raw_name': person,
                    'norm_name': norm,
                    'first_last_key': first_last_key(norm),
                    'is_vacancy': norm == 'OPEN',
                })
    return rows


def extract_pdf_staffing(path):
    reader = PdfReader(path)
    text = '\n'.join(page.extract_text() or '' for page in reader.pages)
    school = infer_school(text, os.path.basename(path))
    rows = []
    current_role = None
    seen = set()
    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        low = line.lower()
        if any(x in low for x in ['please complete this form', 'projected enrollment', 'number of positions', 'teacher’s name', 'next page']):
            continue
        role_guess = normalize_role(line)
        if role_guess != 'Other' and len(normalize_space(line).split()) <= 6:
            current_role = normalize_space(line)
            continue
        if current_role:
            for person in split_names_blob(line):
                norm = normalize_name_for_match(person)
                if not norm:
                    continue
                rec = (school, os.path.basename(path), current_role, normalize_role(current_role), person, norm)
                if rec in seen:
                    continue
                seen.add(rec)
                role = normalize_role(current_role)
                rows.append({
                    'school': school,
                    'source_file': os.path.basename(path),
                    'source_type': 'PDF',
                    'raw_role': current_role,
                    'role': role,
                    'role_family': role_family(role),
                    'raw_name': person,
                    'norm_name': norm,
                    'first_last_key': first_last_key(norm),
                    'is_vacancy': norm == 'OPEN',
                })
    return rows


def extract_staffing_rows(path):
    ext = os.path.splitext(path.lower())[1]
    if ext == '.docx':
        return extract_docx_staffing(path)
    if ext == '.pdf':
        return extract_pdf_staffing(path)
    if ext == '.txt':
        school = infer_school('', os.path.basename(path))
        rows = []
        with open(path, 'r', encoding='utf-8', errors='ignore') as f:
            role = None
            for raw_line in f:
                line = raw_line.strip()
                if not line:
                    continue
                role_guess = normalize_role(line)
                if role_guess != 'Other':
                    role = line
                    continue
                if role:
                    person = clean_name_text(line)
                    norm = normalize_name_for_match(person)
                    if norm:
                        rows.append({
                            'school': school,
                            'source_file': os.path.basename(path),
                            'source_type': 'TXT',
                            'raw_role': role,
                            'role': normalize_role(role),
                            'role_family': role_family(normalize_role(role)),
                            'raw_name': person,
                            'norm_name': norm,
                            'first_last_key': first_last_key(norm),
                            'is_vacancy': norm == 'OPEN',
                        })
        return rows
    return []


def parse_position_master(path):
    df = pd.read_excel(path)
    colmap = {c.lower().strip(): c for c in df.columns}
    required = ['location name', 'job description', 'employee name']
    for req in required:
        if req not in colmap:
            raise ValueError(f'Missing required column in position master: {req}')
    out = pd.DataFrame({
        'school': df[colmap['location name']].astype(str).str.strip(),
        'raw_role': df[colmap['job description']].astype(str).str.strip(),
        'raw_name': df[colmap['employee name']].astype(str).str.strip(),
        'fte': df[colmap['fte']] if 'fte' in colmap else '',
        'emp_num': df[colmap['emp.#']] if 'emp.#' in colmap else '',
    })
    out['school'] = out['school'].map(lambda x: SCHOOL_ALIASES.get(str(x).strip().lower(), x))
    out['display_name'] = out['raw_name'].map(pm_display_name)
    out['role'] = out['raw_role'].map(normalize_role)
    out['role_family'] = out['role'].map(role_family)
    out['norm_name'] = out['raw_name'].map(normalize_name_for_match)
    out['first_last_key'] = out['norm_name'].map(first_last_key)
    out['is_vacancy'] = out['norm_name'].eq('OPEN') | out['raw_name'].str.contains('vacant|open', case=False, na=False)
    return out


def best_fuzzy_match(target_name, candidate_names):
    best_name, best_score = None, -1
    for cand in candidate_names:
        score = max(
            fuzz.token_sort_ratio(target_name, cand),
            fuzz.ratio(target_name, cand),
            fuzz.partial_ratio(target_name, cand),
        )
        if score > best_score:
            best_name, best_score = cand, score
    return best_name, best_score


def run_match(position_master_path, staffing_paths):
    pm = parse_position_master(position_master_path)
    staffing_rows = []
    allocations = []
    for path in staffing_paths:
        staffing_rows.extend(extract_staffing_rows(path))
        if path.lower().endswith('.docx'):
            allocations.extend(extract_allocations_docx(path))
    sf = pd.DataFrame(staffing_rows)
    if sf.empty:
        raise ValueError('No staffing rows could be extracted from the uploaded files.')

    schools_in_scope = sorted(sf['school'].dropna().unique().tolist())
    pm = pm[pm['school'].isin(schools_in_scope)].copy()
    sf = sf[sf['school'].isin(schools_in_scope)].copy()

    matched = []
    review = []
    used_pm = set()

    vacancies_df = sf[sf['is_vacancy']].copy()
    non_vac_sf = sf[(~sf['is_vacancy']) & sf['norm_name'].ne('')].copy()

    # Group PM rows for fast lookups
    exact_key = defaultdict(list)
    first_last_role_key = defaultdict(list)
    family_name_key = defaultdict(list)
    school_name_key = defaultdict(list)
    for idx, row in pm[~pm['is_vacancy']].iterrows():
        exact_key[(row['school'], row['role'], row['norm_name'])].append(idx)
        first_last_role_key[(row['school'], row['role'], row['first_last_key'])].append(idx)
        family_name_key[(row['school'], row['role_family'], row['norm_name'])].append(idx)
        school_name_key[(row['school'], row['norm_name'])].append(idx)

    unmatched_idxs = []
    for idx, row in non_vac_sf.iterrows():
        pm_idx = None
        reason = None
        conf = None

        # Pass 1: exact on school + role + normalized full name
        cands = [i for i in exact_key.get((row['school'], row['role'], row['norm_name']), []) if i not in used_pm]
        if cands:
            pm_idx = cands[0]
            reason = 'Exact match on school, normalized role, and normalized name.'
            conf = 100
        else:
            # Pass 2: exact on school + role + first/last key
            cands = [i for i in first_last_role_key.get((row['school'], row['role'], row['first_last_key']), []) if i not in used_pm]
            if cands:
                pm_idx = cands[0]
                reason = 'Matched on school, normalized role, and first/last name after removing middle initials.'
                conf = 98
            else:
                # Pass 3: exact on school + role family + normalized name
                cands = [i for i in family_name_key.get((row['school'], row['role_family'], row['norm_name']), []) if i not in used_pm]
                if cands:
                    pm_idx = cands[0]
                    reason = 'Matched on exact name within the same school and role family.'
                    conf = 95
                else:
                    # Pass 4: exact on school + normalized name, role mismatch review
                    cands = [i for i in school_name_key.get((row['school'], row['norm_name']), []) if i not in used_pm]
                    if cands:
                        cand_row = pm.loc[cands[0]]
                        review.append({
                            'match_type': 'Review Needed',
                            'confidence': 82,
                            'school': row['school'],
                            'staffing_role': row['role'],
                            'pm_role': cand_row['role'],
                            'staffing_name': row['raw_name'],
                            'pm_candidate_name': cand_row['display_name'],
                            'staffing_file': row['source_file'],
                            'notes': 'Exact name match in the same school, but role does not align. Review needed.',
                        })
                        unmatched_idxs.append(idx)
                        continue
                    else:
                        unmatched_idxs.append(idx)
                        continue

        used_pm.add(pm_idx)
        pm_row = pm.loc[pm_idx]
        matched.append({
            'match_type': 'Matched',
            'confidence': conf,
            'school': row['school'],
            'staffing_role': row['role'],
            'pm_role': pm_row['role'],
            'staffing_name': row['raw_name'],
            'pm_name': pm_row['display_name'],
            'staffing_file': row['source_file'],
            'pm_job_description': pm_row['raw_role'],
            'emp_number': pm_row['emp_num'],
            'fte': pm_row['fte'],
            'notes': reason,
        })

    # Fuzzy review pass within same school and role / role family for still unmatched
    still_unmatched = non_vac_sf.loc[unmatched_idxs].copy()
    final_unmatched = []
    for idx, row in still_unmatched.iterrows():
        sub = pm[(pm['school'] == row['school']) & (~pm.index.isin(used_pm)) & (~pm['is_vacancy'])].copy()
        same_role = sub[sub['role'] == row['role']]
        same_family = sub[sub['role_family'] == row['role_family']]
        pool = same_role if not same_role.empty else same_family
        candidate_names = [n for n in pool['norm_name'].dropna().tolist() if n and n != 'OPEN']
        if candidate_names:
            best_name, score = best_fuzzy_match(row['norm_name'], candidate_names)
            cand_idx = pool[pool['norm_name'] == best_name].index[0]
            cand = pm.loc[cand_idx]
            if score >= 93:
                used_pm.add(cand_idx)
                matched.append({
                    'match_type': 'Fuzzy Auto',
                    'confidence': int(score),
                    'school': row['school'],
                    'staffing_role': row['role'],
                    'pm_role': cand['role'],
                    'staffing_name': row['raw_name'],
                    'pm_name': cand['display_name'],
                    'staffing_file': row['source_file'],
                    'pm_job_description': cand['raw_role'],
                    'emp_number': cand['emp_num'],
                    'fte': cand['fte'],
                    'notes': f'High-confidence fuzzy name match within the same school and role bucket (score {int(score)}).',
                })
                continue
            if score >= 80:
                review.append({
                    'match_type': 'Review Needed',
                    'confidence': int(score),
                    'school': row['school'],
                    'staffing_role': row['role'],
                    'pm_role': cand['role'],
                    'staffing_name': row['raw_name'],
                    'pm_candidate_name': cand['display_name'],
                    'staffing_file': row['source_file'],
                    'notes': f'Possible fuzzy name match inside the same school and role bucket (score {int(score)}).',
                })
        final_unmatched.append(idx)

    matched_df = pd.DataFrame(matched)
    if matched_df.empty:
        matched_df = pd.DataFrame(columns=['match_type','confidence','school','staffing_role','pm_role','staffing_name','pm_name','staffing_file','pm_job_description','emp_number','fte','notes'])
    else:
        matched_df = matched_df.sort_values(['school', 'staffing_role', 'confidence', 'staffing_name'], ascending=[True, True, False, True])

    review_df = pd.DataFrame(review)
    if review_df.empty:
        review_df = pd.DataFrame(columns=['match_type','confidence','school','staffing_role','pm_role','staffing_name','pm_candidate_name','staffing_file','notes'])
    else:
        review_df = review_df.sort_values(['school', 'staffing_role', 'confidence'], ascending=[True, True, False])

    unmatched_staffing_df = non_vac_sf.loc[final_unmatched, ['school', 'role', 'raw_role', 'raw_name', 'source_file']].rename(columns={
        'role': 'staffing_role',
        'raw_role': 'staffing_raw_role',
        'raw_name': 'staffing_name',
    }).sort_values(['school', 'staffing_role', 'staffing_name'])

    unmatched_pm_df = pm[(~pm.index.isin(used_pm)) & (~pm['is_vacancy'])][['school', 'role', 'raw_role', 'display_name', 'fte', 'emp_num']].rename(columns={
        'role': 'pm_role',
        'raw_role': 'pm_raw_role',
        'display_name': 'pm_name',
        'emp_num': 'emp_number',
    }).sort_values(['school', 'pm_role', 'pm_name'])

    vacancies_out = vacancies_df[['school', 'role', 'raw_role', 'raw_name', 'source_file']].rename(columns={
        'role': 'vacancy_role',
        'raw_role': 'vacancy_raw_role',
        'raw_name': 'vacancy_marker',
    }).sort_values(['school', 'vacancy_role'])

    # Role coverage and allocation check
    staffing_counts = non_vac_sf.groupby(['school', 'role']).size().reset_index(name='staffing_filled')
    vacancy_counts = vacancies_df.groupby(['school', 'role']).size().reset_index(name='staffing_vacancies')
    pm_counts = pm[~pm['is_vacancy']].groupby(['school', 'role']).size().reset_index(name='pm_rows')
    coverage = staffing_counts.merge(vacancy_counts, on=['school', 'role'], how='outer').merge(pm_counts, on=['school', 'role'], how='outer').fillna(0)
    for c in ['staffing_filled', 'staffing_vacancies', 'pm_rows']:
        coverage[c] = coverage[c].astype(int)
    coverage['delta_pm_minus_staffing'] = coverage['pm_rows'] - coverage['staffing_filled']
    coverage['status'] = coverage['delta_pm_minus_staffing'].map(lambda x: 'Balanced' if x == 0 else ('PM Higher' if x > 0 else 'Staffing Higher'))

    alloc_df = pd.DataFrame(allocations)
    if not alloc_df.empty:
        alloc_df = alloc_df.groupby(['school', 'role'], as_index=False)['allocation_count'].max()
        allocation_check = alloc_df.merge(staffing_counts, on=['school', 'role'], how='outer').merge(vacancy_counts, on=['school', 'role'], how='outer').merge(pm_counts, on=['school', 'role'], how='outer').fillna(0)
        for c in ['allocation_count', 'staffing_filled', 'staffing_vacancies', 'pm_rows']:
            allocation_check[c] = allocation_check[c].astype(int)
        allocation_check['staffing_gap_vs_allocation'] = allocation_check['staffing_filled'] - allocation_check['allocation_count']
        allocation_check['vacancy_gap'] = allocation_check['allocation_count'] - allocation_check['staffing_filled']
        allocation_check['allocation_status'] = allocation_check.apply(
            lambda r: 'Balanced' if r['staffing_filled'] == r['allocation_count'] else ('Underfilled' if r['staffing_filled'] < r['allocation_count'] else 'Overfilled'),
            axis=1,
        )
        allocation_check = allocation_check.sort_values(['school', 'role'])
    else:
        allocation_check = pd.DataFrame(columns=['school', 'role', 'allocation_count', 'staffing_filled', 'staffing_vacancies', 'pm_rows', 'staffing_gap_vs_allocation', 'vacancy_gap', 'allocation_status'])

    coverage = coverage.sort_values(['school', 'role'])
    coverage['role_sort'] = coverage['role'].map(lambda x: ROLE_ORDER.index(x) if x in ROLE_ORDER else 999)
    coverage = coverage.sort_values(['school', 'role_sort', 'role']).drop(columns=['role_sort'])

    summary = pd.DataFrame([{
        'schools_in_scope': ', '.join(schools_in_scope),
        'staffing_files_loaded': len(staffing_paths),
        'staffing_rows_extracted': len(sf),
        'vacancies_detected': int(vacancies_df.shape[0]),
        'matched_rows': int(matched_df.shape[0]),
        'review_needed_rows': int(review_df.shape[0]),
        'unmatched_staffing_rows': int(unmatched_staffing_df.shape[0]),
        'unmatched_position_master_rows': int(unmatched_pm_df.shape[0]),
    }])

    extracted_staff = sf[['school', 'role', 'raw_role', 'raw_name', 'source_file', 'is_vacancy']].rename(columns={
        'role': 'staffing_role', 'raw_role': 'staffing_raw_role', 'raw_name': 'staffing_name'
    }).sort_values(['school', 'staffing_role', 'staffing_name'])
    extracted_pm = pm[['school', 'role', 'raw_role', 'display_name', 'fte', 'emp_num']].rename(columns={
        'role': 'pm_role', 'raw_role': 'pm_raw_role', 'display_name': 'pm_name', 'emp_num': 'emp_number'
    }).sort_values(['school', 'pm_role', 'pm_name'])

    return {
        'Summary': summary,
        'Matched': matched_df,
        'Review Needed': review_df,
        'Vacancies': vacancies_out,
        'Role Coverage': coverage,
        'Allocation Check': allocation_check,
        'Unmatched Staffing': unmatched_staffing_df,
        'Unmatched Position Master': unmatched_pm_df,
        'Extracted Staffing Rows': extracted_staff,
        'Extracted Position Master Rows': extracted_pm,
    }


def save_output(sheets_dict):
    outname = f"position_master_best_output_{uuid.uuid4().hex[:8]}.xlsx"
    outpath = os.path.join(OUTPUT_FOLDER, outname)
    with pd.ExcelWriter(outpath, engine="openpyxl") as writer:
        for sheet, df in sheets_dict.items():
            sheet_name = sheet[:31]
            df.to_excel(writer, index=False, sheet_name=sheet_name)
            ws = writer.book[sheet_name]
            ws.freeze_panes = "A2"
            for idx, col_name in enumerate(df.columns, start=1):
                sample_vals = [str(col_name)] + [str(v) for v in df[col_name].head(50).tolist()]
                width = min(max(len(v) for v in sample_vals) + 2, 42)
                ws.column_dimensions[ws.cell(row=1, column=idx).column_letter].width = width
    return outpath, outname


@app.route('/', methods=['GET', 'POST'])
def index():
    output_link = None
    if request.method == 'POST':
        try:
            pm_file = request.files.get('pm_file')
            staffing_files = request.files.getlist('staffing_files')
            if not pm_file or not pm_file.filename:
                raise ValueError('Upload the Position Master Excel file.')
            if not allowed(pm_file.filename, ALLOWED_PM):
                raise ValueError('Position Master must be an Excel file (.xlsx or .xls).')
            staffing_files = [f for f in staffing_files if f and f.filename]
            if not staffing_files:
                raise ValueError('Upload at least one staffing file.')
            if len(staffing_files) > MAX_STAFFING_FILES:
                raise ValueError(f'You can upload up to {MAX_STAFFING_FILES} staffing files at a time.')
            for f in staffing_files:
                if not allowed(f.filename, ALLOWED_STAFFING):
                    raise ValueError('Staffing files must be .docx, .pdf, or .txt.')

            batch = uuid.uuid4().hex[:8]
            pm_path = os.path.join(UPLOAD_FOLDER, f'{batch}_{secure_filename(pm_file.filename)}')
            pm_file.save(pm_path)

            staffing_paths = []
            for f in staffing_files:
                path = os.path.join(UPLOAD_FOLDER, f'{batch}_{secure_filename(f.filename)}')
                f.save(path)
                staffing_paths.append(path)

            sheets = run_match(pm_path, staffing_paths)
            _, outname = save_output(sheets)
            output_link = url_for('download_file', filename=outname)
            return render_template('index.html', output_link=output_link, max_files=MAX_STAFFING_FILES)
        except Exception as e:
            flash(str(e), 'error')
            return redirect(url_for('index'))
    return render_template('index.html', output_link=output_link, max_files=MAX_STAFFING_FILES)


@app.route('/download/<filename>')
def download_file(filename):
    return send_file(os.path.join(OUTPUT_FOLDER, filename), as_attachment=True)


if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port, debug=False)
